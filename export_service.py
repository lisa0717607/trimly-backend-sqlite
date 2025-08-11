import os
import json
import tempfile
from datetime import datetime, timedelta
from typing import Dict, Any, Optional, List, Tuple
from pathlib import Path
import asyncio

# 音訊處理
from pydub import AudioSegment
from pydub.utils import which

# 文件生成
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from utils import safe_json_loads, safe_json_dumps, generate_unique_filename

class ExportService:
    """多格式匯出服務"""
    
    def __init__(self):
        # 確保 ffmpeg 可用
        AudioSegment.converter = which("ffmpeg")
        AudioSegment.ffmpeg = which("ffmpeg")
        AudioSegment.ffprobe = which("ffprobe")
        
        # 支援的音訊格式
        self.supported_audio_formats = {
            "mp3": {"quality": "compressed", "extension": ".mp3"},
            "wav": {"quality": "lossless", "extension": ".wav"},
            "flac": {"quality": "lossless", "extension": ".flac"},
            "aac": {"quality": "compressed", "extension": ".aac"},
            "ogg": {"quality": "compressed", "extension": ".ogg"}
        }
        
        # 支援的文字格式
        self.supported_text_formats = {
            "txt": {"description": "純文字檔", "extension": ".txt"},
            "srt": {"description": "字幕檔", "extension": ".srt"},
            "vtt": {"description": "WebVTT 字幕檔", "extension": ".vtt"},
            "docx": {"description": "Word 文件", "extension": ".docx"},
            "json": {"description": "JSON 資料", "extension": ".json"}
        }
    
    async def export_audio(self, audio_file_path: str, output_format: str, 
                          quality_settings: Dict[str, Any] = None) -> Dict[str, Any]:
        """匯出音訊檔案"""
        
        if output_format not in self.supported_audio_formats:
            return {
                "success": False,
                "error": f"Unsupported audio format: {output_format}"
            }
        
        try:
            # 載入音訊檔案
            audio = AudioSegment.from_file(audio_file_path)
            
            # 設定品質參數
            export_params = {}
            
            if output_format == "mp3":
                bitrate = quality_settings.get("bitrate", "192k") if quality_settings else "192k"
                export_params = {
                    "format": "mp3",
                    "bitrate": bitrate,
                    "parameters": ["-q:a", "2"]  # VBR quality
                }
            elif output_format == "wav":
                export_params = {
                    "format": "wav"
                }
            elif output_format == "flac":
                compression_level = quality_settings.get("compression", 5) if quality_settings else 5
                export_params = {
                    "format": "flac",
                    "parameters": ["-compression_level", str(compression_level)]
                }
            elif output_format == "aac":
                bitrate = quality_settings.get("bitrate", "128k") if quality_settings else "128k"
                export_params = {
                    "format": "adts",
                    "codec": "aac",
                    "bitrate": bitrate
                }
            elif output_format == "ogg":
                quality = quality_settings.get("quality", 5) if quality_settings else 5
                export_params = {
                    "format": "ogg",
                    "codec": "libvorbis",
                    "parameters": ["-q:a", str(quality)]
                }
            
            # 生成輸出檔案路徑
            output_filename = generate_unique_filename(
                f"export_{output_format}", 
                self.supported_audio_formats[output_format]["extension"]
            )
            output_path = f"/tmp/{output_filename}"
            
            # 匯出音訊
            audio.export(output_path, **export_params)
            
            # 取得檔案資訊
            file_size = os.path.getsize(output_path)
            duration_ms = len(audio)
            
            return {
                "success": True,
                "output_path": output_path,
                "filename": output_filename,
                "format": output_format,
                "file_size": file_size,
                "duration_ms": duration_ms,
                "quality": self.supported_audio_formats[output_format]["quality"],
                "export_settings": export_params
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"Audio export failed: {str(e)}"
            }
    
    def export_transcript_txt(self, transcript_data: Dict[str, Any], 
                             include_timestamps: bool = False,
                             include_speaker_labels: bool = False) -> Dict[str, Any]:
        """匯出純文字逐字稿"""
        
        try:
            segments = transcript_data.get("segments", [])
            
            lines = []
            
            if include_timestamps and include_speaker_labels:
                # 完整格式：時間戳 + 講者 + 文字
                for segment in segments:
                    start_time = self._format_timestamp(segment.get("start", 0))
                    end_time = self._format_timestamp(segment.get("end", 0))
                    speaker = segment.get("speaker", "Speaker")
                    text = segment.get("text", "").strip()
                    
                    lines.append(f"[{start_time} - {end_time}] {speaker}: {text}")
                    
            elif include_timestamps:
                # 時間戳 + 文字
                for segment in segments:
                    start_time = self._format_timestamp(segment.get("start", 0))
                    text = segment.get("text", "").strip()
                    
                    lines.append(f"[{start_time}] {text}")
                    
            elif include_speaker_labels:
                # 講者 + 文字
                for segment in segments:
                    speaker = segment.get("speaker", "Speaker")
                    text = segment.get("text", "").strip()
                    
                    lines.append(f"{speaker}: {text}")
                    
            else:
                # 純文字
                for segment in segments:
                    text = segment.get("text", "").strip()
                    if text:
                        lines.append(text)
            
            # 合併文字
            content = "\n".join(lines)
            
            # 生成檔案
            filename = generate_unique_filename("transcript", ".txt")
            output_path = f"/tmp/{filename}"
            
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(content)
            
            return {
                "success": True,
                "output_path": output_path,
                "filename": filename,
                "format": "txt",
                "content_preview": content[:200] + "..." if len(content) > 200 else content,
                "total_characters": len(content),
                "total_lines": len(lines)
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"TXT export failed: {str(e)}"
            }
    
    def export_transcript_srt(self, transcript_data: Dict[str, Any]) -> Dict[str, Any]:
        """匯出 SRT 字幕檔"""
        
        try:
            segments = transcript_data.get("segments", [])
            
            srt_lines = []
            
            for i, segment in enumerate(segments, 1):
                start_time = self._format_srt_timestamp(segment.get("start", 0))
                end_time = self._format_srt_timestamp(segment.get("end", 0))
                text = segment.get("text", "").strip()
                
                if text:
                    srt_lines.append(f"{i}")
                    srt_lines.append(f"{start_time} --> {end_time}")
                    srt_lines.append(text)
                    srt_lines.append("")  # 空行分隔
            
            content = "\n".join(srt_lines)
            
            # 生成檔案
            filename = generate_unique_filename("subtitles", ".srt")
            output_path = f"/tmp/{filename}"
            
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(content)
            
            return {
                "success": True,
                "output_path": output_path,
                "filename": filename,
                "format": "srt",
                "subtitle_count": len([s for s in segments if s.get("text", "").strip()]),
                "total_duration": segments[-1].get("end", 0) if segments else 0
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"SRT export failed: {str(e)}"
            }
    
    def export_transcript_vtt(self, transcript_data: Dict[str, Any]) -> Dict[str, Any]:
        """匯出 WebVTT 字幕檔"""
        
        try:
            segments = transcript_data.get("segments", [])
            
            vtt_lines = ["WEBVTT", ""]  # WebVTT 標頭
            
            for segment in segments:
                start_time = self._format_vtt_timestamp(segment.get("start", 0))
                end_time = self._format_vtt_timestamp(segment.get("end", 0))
                text = segment.get("text", "").strip()
                
                if text:
                    vtt_lines.append(f"{start_time} --> {end_time}")
                    vtt_lines.append(text)
                    vtt_lines.append("")  # 空行分隔
            
            content = "\n".join(vtt_lines)
            
            # 生成檔案
            filename = generate_unique_filename("subtitles", ".vtt")
            output_path = f"/tmp/{filename}"
            
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(content)
            
            return {
                "success": True,
                "output_path": output_path,
                "filename": filename,
                "format": "vtt",
                "subtitle_count": len([s for s in segments if s.get("text", "").strip()]),
                "total_duration": segments[-1].get("end", 0) if segments else 0
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"VTT export failed: {str(e)}"
            }
    
    def export_transcript_docx(self, transcript_data: Dict[str, Any], 
                              project_info: Dict[str, Any] = None) -> Dict[str, Any]:
        """匯出 Word 文件"""
        
        try:
            segments = transcript_data.get("segments", [])
            
            # 建立 Word 文件
            doc = Document()
            
            # 添加標題
            title = project_info.get("name", "Transcript") if project_info else "Transcript"
            doc_title = doc.add_heading(title, 0)
            doc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加專案資訊
            if project_info:
                info_para = doc.add_paragraph()
                info_para.add_run("專案資訊").bold = True
                
                doc.add_paragraph(f"專案名稱：{project_info.get('name', 'N/A')}")
                doc.add_paragraph(f"建立時間：{project_info.get('created_at', 'N/A')}")
                doc.add_paragraph(f"總時長：{self._format_duration(project_info.get('duration', 0))}")
                doc.add_paragraph(f"匯出時間：{datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')}")
                
                doc.add_paragraph()  # 空行
            
            # 添加逐字稿內容
            transcript_heading = doc.add_heading("逐字稿內容", level=1)
            
            for i, segment in enumerate(segments, 1):
                start_time = self._format_timestamp(segment.get("start", 0))
                end_time = self._format_timestamp(segment.get("end", 0))
                speaker = segment.get("speaker", "Speaker")
                text = segment.get("text", "").strip()
                
                if text:
                    # 時間戳和講者
                    time_para = doc.add_paragraph()
                    time_run = time_para.add_run(f"[{start_time} - {end_time}] {speaker}:")
                    time_run.bold = True
                    
                    # 文字內容
                    text_para = doc.add_paragraph(text)
                    text_para.paragraph_format.left_indent = Inches(0.5)
                    
                    # 段落間距
                    if i < len(segments):
                        doc.add_paragraph()
            
            # 生成檔案
            filename = generate_unique_filename("transcript", ".docx")
            output_path = f"/tmp/{filename}"
            
            doc.save(output_path)
            
            return {
                "success": True,
                "output_path": output_path,
                "filename": filename,
                "format": "docx",
                "segment_count": len(segments),
                "total_pages": len(doc.paragraphs) // 20 + 1  # 估算頁數
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"DOCX export failed: {str(e)}"
            }
    
    def export_transcript_json(self, transcript_data: Dict[str, Any], 
                              include_metadata: bool = True) -> Dict[str, Any]:
        """匯出 JSON 格式的完整資料"""
        
        try:
            export_data = {
                "transcript": transcript_data
            }
            
            if include_metadata:
                export_data["metadata"] = {
                    "export_date": datetime.utcnow().isoformat(),
                    "format_version": "1.0",
                    "total_segments": len(transcript_data.get("segments", [])),
                    "total_duration": transcript_data.get("duration", 0),
                    "language": transcript_data.get("language", "unknown")
                }
            
            # 生成檔案
            filename = generate_unique_filename("transcript_data", ".json")
            output_path = f"/tmp/{filename}"
            
            with open(output_path, "w", encoding="utf-8") as f:
                json.dump(export_data, f, ensure_ascii=False, indent=2)
            
            return {
                "success": True,
                "output_path": output_path,
                "filename": filename,
                "format": "json",
                "data_size": len(json.dumps(export_data)),
                "includes_metadata": include_metadata
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"JSON export failed: {str(e)}"
            }
    
    def create_export_package(self, exports: List[Dict[str, Any]], 
                             package_name: str = None) -> Dict[str, Any]:
        """建立匯出套件（ZIP 檔案）"""
        
        try:
            import zipfile
            
            if not package_name:
                package_name = f"trimly_export_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}"
            
            zip_filename = f"{package_name}.zip"
            zip_path = f"/tmp/{zip_filename}"
            
            with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                for export in exports:
                    if export.get("success") and export.get("output_path"):
                        file_path = export["output_path"]
                        if os.path.exists(file_path):
                            # 使用原始檔名加入 ZIP
                            arcname = export.get("filename", os.path.basename(file_path))
                            zipf.write(file_path, arcname)
            
            # 取得 ZIP 檔案大小
            zip_size = os.path.getsize(zip_path)
            
            return {
                "success": True,
                "package_path": zip_path,
                "package_filename": zip_filename,
                "package_size": zip_size,
                "included_files": len(exports),
                "format": "zip"
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"Package creation failed: {str(e)}"
            }
    
    def _format_timestamp(self, seconds: float) -> str:
        """格式化時間戳（HH:MM:SS）"""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        return f"{hours:02d}:{minutes:02d}:{secs:02d}"
    
    def _format_srt_timestamp(self, seconds: float) -> str:
        """格式化 SRT 時間戳（HH:MM:SS,mmm）"""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        milliseconds = int((seconds % 1) * 1000)
        return f"{hours:02d}:{minutes:02d}:{secs:02d},{milliseconds:03d}"
    
    def _format_vtt_timestamp(self, seconds: float) -> str:
        """格式化 WebVTT 時間戳（HH:MM:SS.mmm）"""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        milliseconds = int((seconds % 1) * 1000)
        return f"{hours:02d}:{minutes:02d}:{secs:02d}.{milliseconds:03d}"
    
    def _format_duration(self, seconds: float) -> str:
        """格式化時長"""
        if seconds < 60:
            return f"{seconds:.1f} 秒"
        elif seconds < 3600:
            minutes = seconds / 60
            return f"{minutes:.1f} 分鐘"
        else:
            hours = seconds / 3600
            return f"{hours:.1f} 小時"
    
    def get_supported_formats(self) -> Dict[str, Any]:
        """取得支援的匯出格式"""
        
        return {
            "audio_formats": self.supported_audio_formats,
            "text_formats": self.supported_text_formats,
            "package_formats": {
                "zip": {"description": "ZIP 壓縮檔", "extension": ".zip"}
            }
        }
    
    def cleanup_temp_files(self, file_paths: List[str]):
        """清理暫存檔案"""
        
        for file_path in file_paths:
            try:
                if os.path.exists(file_path):
                    os.remove(file_path)
            except Exception as e:
                print(f"Failed to cleanup file {file_path}: {str(e)}")

# 全域匯出服務實例
export_service = ExportService()

